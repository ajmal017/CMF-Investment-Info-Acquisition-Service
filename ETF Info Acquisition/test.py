from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from enum import Enum

descriptionInfo = dict()
descriptionInfo["Two Line Description"] = "Hello World"

nameOfETF = "VTI"

holdingsData = [{'Ticker': 'BRK-B', 'Weight': '1.41%', 'Market Capitalization': '503.47B', 'Trailing PE': '18.80', 'Forward PE': '18.56', 'PEG Ratio (5 yr expected)': '0.86', 'Price/Sales (ttm)': '2.01', 'Price/Book (mrq)': '0.00', 'Enterprise Value/Revenue': '-0.05', 'Enterprise Value/EBITDA': '-0.26', 'Profit Margin': '10.70%', 'ROA (ttm)': '3.42%', 'ROE (ttm)': '7.50%', 'Quarterly Revenue Growth (yoy)': '5.50%', 'EBITDA': '49.25B', 'Quarterly Earnings Growth (yoy)': 'N/A', 'Total ESG Score': '45', 'Total ESG Percentile': '9th percentile', 'Environment Score': '41', 'Environment Percentile': '43rd percentile', 'Social Score': '49', 'Social Percentile': '0th percentile', 'Governmental Score': '44', 'Governmental Percentile': '14th percentile', 'Sector': 'Financial Services', 'Industry': 'Insurance - Diversified', 'EBITDA Margin': 'Unavailable', 'Gross Profit Margin': 'Unavailable', 'Net Profit Margin': 'Unavailable'}, {'Ticker': 'BRK-B', 'Weight': '1.41%', 'Market Capitalization': '503.47B', 'Trailing PE': '18.80', 'Forward PE': '18.56', 'PEG Ratio (5 yr expected)': '0.86', 'Price/Sales (ttm)': '2.01', 'Price/Book (mrq)': '0.00', 'Enterprise Value/Revenue': '-0.05', 'Enterprise Value/EBITDA': '-0.26', 'Profit Margin': '10.70%', 'ROA (ttm)': '3.42%', 'ROE (ttm)': '7.50%', 'Quarterly Revenue Growth (yoy)': '5.50%', 'EBITDA': '49.25B', 'Quarterly Earnings Growth (yoy)': 'N/A', 'Total ESG Score': '45', 'Total ESG Percentile': '9th percentile', 'Environment Score': '41', 'Environment Percentile': '43rd percentile', 'Social Score': '49', 'Social Percentile': '0th percentile', 'Governmental Score': '44', 'Governmental Percentile': '14th percentile', 'Sector': 'Financial Services', 'Industry': 'Insurance - Diversified', 'EBITDA Margin': 'Unavailable', 'Gross Profit Margin': 'Unavailable', 'Net Profit Margin': 'Unavailable'}, {'Ticker': 'BRK-B', 'Weight': '1.41%', 'Market Capitalization': '503.47B', 'Trailing PE': '18.80', 'Forward PE': '18.56', 'PEG Ratio (5 yr expected)': '0.86', 'Price/Sales (ttm)': '2.01', 'Price/Book (mrq)': '0.00', 'Enterprise Value/Revenue': '-0.05', 'Enterprise Value/EBITDA': '-0.26', 'Profit Margin': '10.70%', 'ROA (ttm)': '3.42%', 'ROE (ttm)': '7.50%', 'Quarterly Revenue Growth (yoy)': '5.50%', 'EBITDA': '49.25B', 'Quarterly Earnings Growth (yoy)': 'N/A', 'Total ESG Score': '45', 'Total ESG Percentile': '9th percentile', 'Environment Score': '41', 'Environment Percentile': '43rd percentile', 'Social Score': '49', 'Social Percentile': '0th percentile', 'Governmental Score': '44', 'Governmental Percentile': '14th percentile', 'Sector': 'Financial Services', 'Industry': 'Insurance - Diversified', 'EBITDA Margin': 'Unavailable', 'Gross Profit Margin': 'Unavailable', 'Net Profit Margin': 'Unavailable'}, {'Ticker': 'BRK-B', 'Weight': '1.41%', 'Market Capitalization': '503.47B', 'Trailing PE': '18.80', 'Forward PE': '18.56', 'PEG Ratio (5 yr expected)': '0.86', 'Price/Sales (ttm)': '2.01', 'Price/Book (mrq)': '0.00', 'Enterprise Value/Revenue': '-0.05', 'Enterprise Value/EBITDA': '-0.26', 'Profit Margin': '10.70%', 'ROA (ttm)': '3.42%', 'ROE (ttm)': '7.50%', 'Quarterly Revenue Growth (yoy)': '5.50%', 'EBITDA': '49.25B', 'Quarterly Earnings Growth (yoy)': 'N/A', 'Total ESG Score': '45', 'Total ESG Percentile': '9th percentile', 'Environment Score': '41', 'Environment Percentile': '43rd percentile', 'Social Score': '49', 'Social Percentile': '0th percentile', 'Governmental Score': '44', 'Governmental Percentile': '14th percentile', 'Sector': 'Financial Services', 'Industry': 'Insurance - Diversified', 'EBITDA Margin': 'Unavailable', 'Gross Profit Margin': 'Unavailable', 'Net Profit Margin': 'Unavailable'}, {'Ticker': 'BRK-B', 'Weight': '1.41%', 'Market Capitalization': '503.47B', 'Trailing PE': '18.80', 'Forward PE': '18.56', 'PEG Ratio (5 yr expected)': '0.86', 'Price/Sales (ttm)': '2.01', 'Price/Book (mrq)': '0.00', 'Enterprise Value/Revenue': '-0.05', 'Enterprise Value/EBITDA': '-0.26', 'Profit Margin': '10.70%', 'ROA (ttm)': '3.42%', 'ROE (ttm)': '7.50%', 'Quarterly Revenue Growth (yoy)': '5.50%', 'EBITDA': '49.25B', 'Quarterly Earnings Growth (yoy)': 'N/A', 'Total ESG Score': '45', 'Total ESG Percentile': '9th percentile', 'Environment Score': '41', 'Environment Percentile': '43rd percentile', 'Social Score': '49', 'Social Percentile': '0th percentile', 'Governmental Score': '44', 'Governmental Percentile': '14th percentile', 'Sector': 'Financial Services', 'Industry': 'Insurance - Diversified', 'EBITDA Margin': 'Unavailable', 'Gross Profit Margin': 'Unavailable', 'Net Profit Margin': 'Unavailable'}]



distributionInfo = {'Asset Allocation': ['Common equity (100%)', 'ETF Cash Component (0%)'], 'Sector Breakdown': ['Technology (20%)', 'Healthcare (14%)', 'Financials (13%)', 'Industrials (10%)', 'Consumer, Cyclical (10%)', 'Communications (9%)', 'Consumer, Non-Cyclical (7%)', 'Energy (4%)', 'Real Estate (4%)', 'Utilities (3%)', 'Other (3%)', 'Basic Materials (2%)', 'ETF Cash Component (0%)'], 'Market Cap Breakdown': ['Large cap (79%)', 'Mid cap (13%)', 'Small cap (6%)', 'Micro cap (2%)', 'Unknown (0%)', 'ETF Cash Component (0%)'], 'Region Breakdown': ['North America (98%)', 'Europe (2%)', 'Latin America (0%)', 'Asia (0%)', 'Asia-Pacific (0%)'], 'Market Tier Breakdown': ['United States (98%)', 'Developed Markets (ex-US) (2%)', 'Emerging Markets (0%)', 'Not-designated (0%)'], 'Country Breakdown': ['United States (98%)', 'United Kingdom (1%)', 'Ireland (1%)', 'Switzerland (0%)', 'Canada (0%)', 'Germany (0%)', 'Mexico (0%)', 'Spain (0%)', 'Netherlands (0%)', 'France (0%)', 'China (0%)', 'Brazil (0%)', 'Japan (0%)', 'Thailand (0%)', 'Luxembourg (0%)', 'Puerto Rico (0%)']}


expenseInfo = {'ETF Average': '0.545%', 'ETF Wtd. Average': '0.20%', 'Category Average': '0.38%', 'Expense Ratio': '0.03%'}

esgScoresInfo = {'MSCI ESG Rating': 'BBB', 'MSCI ESG Quality Score': '5.18 / 10', 'Peer Group Percentile Rank': '38.34', 'Global Percentile Rank': '43.47', 'SRI Screening Criteria Exposure': '11.10%', 'Exposure to Sustainable Impact Solutions': '5.72%', 'Weighted Average Carbon Intensity (t CO2e/$M Sales)': '190.86'}

performanceInfo = {'1 Week Return': {'Peer Ranking': '122 of 214', 'Value': '2.01%'}, '4 Week Return': {'Peer Ranking': '123 of 210', 'Value': '-0.94%'}, '13 Week Return': {'Peer Ranking': '119 of 204', 'Value': '3.52%'}, '26 Week Return': {'Peer Ranking': '125 of 196', 'Value': '6.48%'}, 'Year to Date Return': {'Peer Ranking': '97 of 194', 'Value': '14.59%'}, '1 Year Return': {'Peer Ranking': '105 of 176', 'Value': '3.35%'}, '3 Year Return': {'Peer Ranking': '64 of 122', 'Value': '43.16%'}, '5 Year Return': {'Peer Ranking': '52 of 93', 'Value': '59.10%', 'P/E': '23.91', 'P/B': '3.26'}}

riskInfo = {'5 Day Volatility': {'Peer Ranking': '75of 214', 'Value': '108.28%'}, '20 Day Volatility': {'Peer Ranking': '95of 209', 'Value': '16.47%'}, '50 Day Volatility': {'Peer Ranking': '96of 206', 'Value': '12.61%'}, '200 Day Volatility': {'Peer Ranking': '94of 194', 'Value': '12.07%'}, 'Beta': {'Peer Ranking': '78of 207', 'Value': '1.02'}, 'Standard Deviation': {'Peer Ranking': '23of 175', 'Value': '5.96%'}}

valuationRatios = {'P/E Ratio (TTM)': {'Company': '19.63', 'Industry': '6.57', 'Sector': '20.83'}, 'P/E Low - Last 5 Yrs.': {'Company': '3.94', 'Industry': '2.90', 'Sector': '13.62'}, 'Beta': {'Company': '1.04', 'Industry': '0.98', 'Sector': '1.32'}, 'Price to Sales (TTM)': {'Company': '104.40', 'Industry': '33.57', 'Sector': '10.07'}, 'Price to Tangible Book (MRQ)': {'Company': '1.96', 'Industry': '1.96', 'Sector': '3.38'}}

dividendsRatios = {'Dividend Yield': {'Company': '1.93', 'Industry': '1.87', 'Sector': '2.43'}, 'Dividend 5 Year Growth Rate': {'Company': '11.04', 'Industry': '6.99', 'Sector': '17.93'}, 'Payout Ratio(TTM)': {'Company': '25.16', 'Industry': '8.19', 'Sector': '37.61'}}

growthRatesRatios = {'Sales (MRQ) vs Qtr. 1 Yr. Ago': {'Company': '39.49', 'Industry': '390.14', 'Sector': '21.21'}, 'Sales - 5 Yr. Growth Rate': {'Company': '72.01', 'Industry': '58.28', 'Sector': '11.40'}, 'EPS (MRQ) vs Qtr. 1 Yr. Ago': {'Company': '-69.79', 'Industry': '268.35', 'Sector': '19.99'}, 'EPS - 5 Yr. Growth Rate': {'Company': '-1.22', 'Industry': '34.24', 'Sector': '9.52'}, 'Capital Spending - 5 Yr. Growth Rate': {'Company': '--', 'Industry': '0.00', 'Sector': '1.00'}}


financialStrengthsRatios = {'Quick Ratio (MRQ)': {'Company': '--', 'Industry': '0.60', 'Sector': '0.04'}, 'LT Debt to Equity (MRQ)': {'Company': '0.00', 'Industry': '0.04', 'Sector': '32.07'}, 'Interest Coverage (TTM)': {'Company': '--', 'Industry': '0.01', 'Sector': '15.63'}}


profitabilityRatios = {'Gross Margin (TTM)': {'Company': '--', 'Industry': '84.82', 'Sector': '2.68'}, 'EBITD - 5 Yr. Avg': {'Company': '97.91', 'Industry': '69.68', 'Sector': '47.35'}, 'Operating Margin (TTM)': {'Company': '98.02', 'Industry': '91.43', 'Sector': '41.08'}, 'Pre-Tax Margin - 5 Yr. Avg.': {'Company': '774.69', 'Industry': '158.99', 'Sector': '40.43'}, 'Net Profit Margin (TTM)': {'Company': '361.93', 'Industry': '105.46', 'Sector': '32.42'}, 'Effective Tax Rate - 5 Yr. Avg.': {'Company': '--', 'Industry': '0.01', 'Sector': '21.25'}}


efficiencyRatios = {'Revenue/Employee (TTM)': {'Company': '--', 'Industry': '0', 'Sector': '30,518,125'}, 'Inventory Turnover (TTM)': {'Company': '--', 'Industry': '0.00', 'Sector': '0.90'}}


managementEffectivenessRatios = {'Return on Assets (TTM)': {'Company': '7.64', 'Industry': '18.71', 'Sector': '2.82'}, 'Return on Investment - 5 Yr. Avg.': {'Company': '17.26', 'Industry': '12.99', 'Sector': '0.36'}, 'Return on Equity (TTM)': {'Company': '7.64', 'Industry': '19.42', 'Sector': '16.22'}}


liquidityRating = "AAA"






class TableTitleStyle(Enum):
    SIDE = 1
    TOP = 2
    BOTH = 3


document = Document()
normalStyle = document.styles['Normal']
font = normalStyle.font
font.name = 'Times New Roman'
font.size = Pt(12)



# =============== Useful Functions ================
def addTitle(name):
    t1 = document.add_paragraph()
    t1.add_run(name).bold = True
    t1.style = normalStyle

def addTable(data, titleFormat=TableTitleStyle.TOP):
    # Data should be a 2 dimensional matrix
    nOfRows = len(data)
    nOfCols = len(data[0])
    table = document.add_table(rows=nOfRows, cols=nOfCols)
    table.style = 'TableGrid'

    for indexr, eachRow in enumerate(table.rows):
        for indexc, cell in enumerate(eachRow.cells):
            cell.text = data[indexr][indexc]

    if titleFormat == TableTitleStyle.TOP:
        for eachCell in table.rows[0].cells:
            run = eachCell.paragraphs[0].runs[0]
            run.font.bold = True
    elif titleFormat == TableTitleStyle.SIDE:
        for eachCell in table.columns[0].cells:
            run = eachCell.paragraphs[0].runs[0]
            run.font.bold = True
    else:
        for eachCell in table.rows[0].cells:
            run = eachCell.paragraphs[0].runs[0]
            run.font.bold = True

        for eachCell in table.columns[0].cells:
            run = eachCell.paragraphs[0].runs[0]
            run.font.bold = True

def addParagraph(paragraphText, alignment = WD_ALIGN_PARAGRAPH.LEFT):
    p = document.add_paragraph(paragraphText)
    p.style = normalStyle
    p.alignment = alignment

def newLine():
    addParagraph("")

def transpose(twoDArray):
    newArray = list()
    nC = len(twoDArray)
    nR = len(twoDArray[0])
    for i in range(nR):
        innerList = [None] * nC
        newArray.append(innerList)

    for indexi, elementi in enumerate(twoDArray):
        for indexj, elementj in enumerate(elementi):
            newArray[indexj][indexi] = elementj

    return newArray
# =========== Title ==========
ticker = "VTI"
titleOfDoc = document.add_heading()
titleOfDoc.add_run(ticker + " Investment Proposal").bold = True
titleOfDoc.style = normalStyle
titleOfDoc.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============ Basic Information ==========
addTitle("Basic Information")

value = [["Name", nameOfETF], ["Ticker", ticker], ["Brief Introduction", descriptionInfo["Two Line Description"]], ["Proposed Investment Duration", "FILL THIS IN"], ["Proposed Number of Shares to Purchase", "FILL THIS IN"]]
addTable(value, titleFormat=TableTitleStyle.TOP)
newLine()

# =================== Graph ==================
addTitle("Historical Graph and Analysis")
addParagraph("Price Trend Over Past 5 Years", alignment=WD_ALIGN_PARAGRAPH.CENTER)
addParagraph("SCREENSHOT A GRAPH HERE", alignment=WD_ALIGN_PARAGRAPH.CENTER)
newLine()
addParagraph("Historical Performance", alignment=WD_ALIGN_PARAGRAPH.CENTER)
titles = list(performanceInfo)
peers = ["Peer Ranking"]
value = ["Value"]
for eachItem in titles:
    entries = performanceInfo[eachItem]
    peers.append(entries["Peer Ranking"])
    value.append(entries["Value"])

performanceArray = transpose([["Name"] + titles, peers, value])
addTable(performanceArray, titleFormat=TableTitleStyle.BOTH)
newLine()

# ==================== Behaviors ===============
addTitle("Behaviour during Bull and Bear Markets")
newLine()
addTitle("How the ETF reacted to major economic changes (e.g. trade war)")
newLine()

# ================= Risk Analysis ================
addTitle("Risk Analysis")
addParagraph("Historical Volatility", alignment=WD_ALIGN_PARAGRAPH.CENTER)
titles = list(riskInfo)
peers = ["Peer Ranking"]
value = ["Value"]
for eachItem in titles:
    entries = riskInfo[eachItem]
    peers.append(entries["Peer Ranking"])
    value.append(entries["Value"])

riskInfoArray = transpose([["Name"] + titles, peers, value])
addTable(riskInfoArray, titleFormat=TableTitleStyle.BOTH)

addParagraph("DO SOME ANALYSIS HERE")
newLine()

# ================= Valuation ================
addTitle("Investment Rationale")
addParagraph("Part I - Valuation", alignment=WD_ALIGN_PARAGRAPH.CENTER)
titles = list(valuationRatios)
company = ["Company"]
industry = ["Industry"]
sector = ["Sector"]
for eachItem in titles:
    entries = valuationRatios[eachItem]
    company.append(entries["Company"])
    industry.append(entries["Industry"])
    sector.append(entries["Sector"])

valuationArray = transpose([["Titles"] + titles, company, industry, sector])
addTable(valuationArray, titleFormat=TableTitleStyle.BOTH)
addParagraph("ADD SOME ANALYSIS HERE")
newLine()

# ================= Dividends ================
addParagraph("Part II - Dividends", alignment=WD_ALIGN_PARAGRAPH.CENTER)
titles = list(dividendsRatios)
company = ["Company"]
industry = ["Industry"]
sector = ["Sector"]
for eachItem in titles:
    entries = dividendsRatios[eachItem]
    company.append(entries["Company"])
    industry.append(entries["Industry"])
    sector.append(entries["Sector"])

dividendsArray = transpose([["Titles"] + titles, company, industry, sector])
addTable(dividendsArray, titleFormat=TableTitleStyle.BOTH)
addParagraph("ADD SOME ANALYSIS HERE")
newLine()

# ================= Profitability ================
addParagraph("Part III - Profitability", alignment=WD_ALIGN_PARAGRAPH.CENTER)
titles = list(profitabilityRatios)
company = ["Company"]
industry = ["Industry"]
sector = ["Sector"]
for eachItem in titles:
    entries = profitabilityRatios[eachItem]
    company.append(entries["Company"])
    industry.append(entries["Industry"])
    sector.append(entries["Sector"])

profitabilityArray = transpose([["Titles"] + titles, company, industry, sector])
addTable(profitabilityArray, titleFormat=TableTitleStyle.BOTH)
newLine()

# ================= Growth Rate Analysis  ================
addParagraph("Part IV - Growth Rate Analysis", alignment=WD_ALIGN_PARAGRAPH.CENTER)
titles = list(growthRatesRatios)
company = ["Company"]
industry = ["Industry"]
sector = ["Sector"]
for eachItem in titles:
    entries = growthRatesRatios[eachItem]
    company.append(entries["Company"])
    industry.append(entries["Industry"])
    sector.append(entries["Sector"])

growthRateArray = transpose([["Titles"] + titles, company, industry, sector])
addTable(growthRateArray, titleFormat=TableTitleStyle.BOTH)
newLine()


# ================= Management Efficiency ================
addParagraph("Part V - Management Efficiency", alignment=WD_ALIGN_PARAGRAPH.CENTER)
titles = list(managementEffectivenessRatios)
company = ["Company"]
industry = ["Industry"]
sector = ["Sector"]
for eachItem in titles:
    entries = managementEffectivenessRatios[eachItem]
    company.append(entries["Company"])
    industry.append(entries["Industry"])
    sector.append(entries["Sector"])

managementEffectivenessArray = transpose([["Titles"] + titles, company, industry, sector])
addTable(managementEffectivenessArray, titleFormat=TableTitleStyle.BOTH)
newLine()

# ================ Miscellaneous Information ===============
addParagraph("Part VI - Miscellaneous Information", alignment=WD_ALIGN_PARAGRAPH.CENTER)
titles = list(expenseInfo)
values = ["Value"]
for eachItem in titles:
    values.append(expenseInfo[eachItem])

titles.append("Liquidity")
values.append(liquidityRating)
miscellanousArray = transpose([["Title"] + titles, values])

addTable(miscellanousArray, titleFormat=TableTitleStyle.BOTH)
newLine()

# ================= Top Five Holdings ================
addTitle("Top Five Holdings")
titles = list(holdingsData[0])
ultimateArray = [titles]
for eachCompany in holdingsData:
    companyInfo = []
    for eachTitle in titles:
        companyInfo.append(eachCompany[eachTitle])

    ultimateArray.append(companyInfo)

ultimateArray = transpose(ultimateArray)
addTable(ultimateArray, titleFormat=TableTitleStyle.SIDE)
newLine()

# ================= ESG Analysis ================
addTitle("ESG Analysis")
keys = list(esgScoresInfo)
values = list(esgScoresInfo.values())
esgArray = transpose([keys, values])
addTable(esgArray, titleFormat=TableTitleStyle.SIDE)

newLine()

# ================= Summary ================
addTitle("Summary")
addParagraph("ADD SUMMARY HERE")
newLine()

document.save('test.docx')
